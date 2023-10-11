from django.contrib import messages
from django.shortcuts import render, redirect
from .forms import UserProfileForm
from docx import Document
from django.http import HttpResponse
from django.core.mail import send_mail
import math, random

def registration_view(request):
    if request.method == 'POST':
        form = UserProfileForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()

            doc = Document()
            user_profile = form.instance
            doc.add_paragraph(f"Full Name: {user_profile.full_name}")
            doc.add_paragraph(f"Email: {user_profile.email}")
            doc.add_paragraph(f"Profile Image: {user_profile.profile_image}")
            doc.add_paragraph(f"Address: {user_profile.address}")
            doc.add_paragraph(f"Phone Number: {user_profile.phone_number}")
            doc.add_paragraph(f"Mode: {user_profile.get_mode_display()}")
            doc.add_paragraph(f"Skills: {user_profile.skills}")
            doc.add_paragraph(f"Subjects: {user_profile.get_subjects_display()}")
            doc.add_paragraph(f"Role: {user_profile.get_role_display()}")
            doc.add_paragraph(f"Class: {user_profile.get_select_class_display()}")  # Display selected class
            doc.save(f'user_profile_{user_profile.id}.docx')

            messages.success(request, 'Your record has been submitted. We will inform you after verification.')
            return redirect('result')
    else:
        form = UserProfileForm()
    return render(request, 'registration/registration_success.html', {'form': form})
def registration_success(request):
    return render(request, 'registration/registration_form.html')

def generateOTP():
    digits = "0123456789"
    OTP = ""
    for i in range(4):
        OTP += digits[math.floor(random.random() * 10)]
    return OTP

def send_otp(request):
    email = request.POST.get("email")  # Use POST instead of GET
    print(email)
    o = generateOTP()
    htmlgen = f'<p>Your OTP is <strong>{o}</strong></p>'
    send_mail('OTP request', o, 'neerajgupta5963@gmail.com', [email], fail_silently=False, html_message=htmlgen)
    return HttpResponse(o)
def result(request):
    return render(request, 'registration/result.html')
def header_view(request):
    return render(request, 'header.html')